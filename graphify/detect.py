# file discovery, type classification, and corpus health checks
from __future__ import annotations
import datetime
import fnmatch
import hashlib
import json
import os
import re
import sys
from enum import Enum
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from graphify.output import ResolvedOutput

from graphify.corpus_prune import (
    _OUTPUT_MANIFEST_NAME,
    _OUTPUT_MANIFEST_VERSION,
    _load_output_manifest,
    build_prior_files,
    dir_prune_reason,
)


class FileType(str, Enum):
    CODE = "code"
    DOCUMENT = "document"
    PAPER = "paper"
    IMAGE = "image"
    VIDEO = "video"


_MANIFEST_PATH = "graphify-out/manifest.json"

CODE_EXTENSIONS = {'.py', '.ts', '.js', '.jsx', '.tsx', '.go', '.rs', '.java', '.cpp', '.cc', '.cxx', '.c', '.h', '.hpp', '.rb', '.swift', '.kt', '.kts', '.cs', '.scala', '.php', '.lua', '.toc', '.zig', '.ps1', '.ex', '.exs', '.m', '.mm', '.jl', '.vue', '.svelte'}
DOC_EXTENSIONS = {'.md', '.txt', '.rst'}
PAPER_EXTENSIONS = {'.pdf'}
IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.svg'}
OFFICE_EXTENSIONS = {'.docx', '.xlsx'}
VIDEO_EXTENSIONS = {'.mp4', '.mov', '.webm', '.mkv', '.avi', '.m4v', '.mp3', '.wav', '.m4a', '.ogg'}

CORPUS_WARN_THRESHOLD = 50_000    # words - below this, warn "you may not need a graph"
CORPUS_UPPER_THRESHOLD = 500_000  # words - above this, warn about token cost
FILE_COUNT_UPPER = 200             # files - above this, warn about token cost

# Files that may contain secrets - skip silently
_SENSITIVE_PATTERNS = [
    re.compile(r'(^|[\\/])\.(env|envrc)(\.|$)', re.IGNORECASE),
    re.compile(r'\.(pem|key|p12|pfx|cert|crt|der|p8)$', re.IGNORECASE),
    re.compile(r'(credential|secret|passwd|password|token|private_key)', re.IGNORECASE),
    re.compile(r'(id_rsa|id_dsa|id_ecdsa|id_ed25519)(\.pub)?$'),
    re.compile(r'(\.netrc|\.pgpass|\.htpasswd)$', re.IGNORECASE),
    re.compile(r'(aws_credentials|gcloud_credentials|service.account)', re.IGNORECASE),
]

# Signals that a .md/.txt file is actually a converted academic paper
_PAPER_SIGNALS = [
    re.compile(r'\barxiv\b', re.IGNORECASE),
    re.compile(r'\bdoi\s*:', re.IGNORECASE),
    re.compile(r'\babstract\b', re.IGNORECASE),
    re.compile(r'\bproceedings\b', re.IGNORECASE),
    re.compile(r'\bjournal\b', re.IGNORECASE),
    re.compile(r'\bpreprint\b', re.IGNORECASE),
    re.compile(r'\\cite\{'),          # LaTeX citation
    re.compile(r'\[\d+\]'),           # Numbered citation [1], [23] (inline)
    re.compile(r'\[\n\d+\n\]'),       # Numbered citation spread across lines (markdown conversion)
    re.compile(r'eq\.\s*\d+|equation\s+\d+', re.IGNORECASE),
    re.compile(r'\d{4}\.\d{4,5}'),   # arXiv ID like 1706.03762
    re.compile(r'\bwe propose\b', re.IGNORECASE),   # common academic phrasing
    re.compile(r'\bliterature\b', re.IGNORECASE),   # "from the literature"
]
_PAPER_SIGNAL_THRESHOLD = 3  # need at least this many signals to call it a paper


def _is_sensitive(path: Path) -> bool:
    """Return True if this file likely contains secrets and should be skipped."""
    name = path.name
    full = str(path)
    return any(p.search(name) or p.search(full) for p in _SENSITIVE_PATTERNS)


def _looks_like_paper(path: Path) -> bool:
    """Heuristic: does this text file read like an academic paper?"""
    try:
        # Only scan first 3000 chars for speed
        text = path.read_text(encoding="utf-8", errors="ignore")[:3000]
        hits = sum(1 for pattern in _PAPER_SIGNALS if pattern.search(text))
        return hits >= _PAPER_SIGNAL_THRESHOLD
    except Exception:
        return False


_ASSET_DIR_MARKERS = {".imageset", ".xcassets", ".appiconset", ".colorset", ".launchimage"}


def classify_file(path: Path) -> FileType | None:
    # Compound extensions must be checked before simple suffix lookup
    if path.name.lower().endswith(".blade.php"):
        return FileType.CODE
    ext = path.suffix.lower()
    if ext in CODE_EXTENSIONS:
        return FileType.CODE
    if ext in PAPER_EXTENSIONS:
        # PDFs inside Xcode asset catalogs are vector icons, not papers
        if any(part.endswith(tuple(_ASSET_DIR_MARKERS)) for part in path.parts):
            return None
        return FileType.PAPER
    if ext in IMAGE_EXTENSIONS:
        return FileType.IMAGE
    if ext in DOC_EXTENSIONS:
        # Check if it's a converted paper
        if _looks_like_paper(path):
            return FileType.PAPER
        return FileType.DOCUMENT
    if ext in OFFICE_EXTENSIONS:
        return FileType.DOCUMENT
    if ext in VIDEO_EXTENSIONS:
        return FileType.VIDEO
    return None


def extract_pdf_text(path: Path) -> str:
    """Extract plain text from a PDF file using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)
    except Exception:
        return ""


def docx_to_markdown(path: Path) -> str:
    """Convert a .docx file to markdown text using python-docx."""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(str(path))
        lines = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            if style.startswith("Heading 1"):
                lines.append(f"# {text}")
            elif style.startswith("Heading 2"):
                lines.append(f"## {text}")
            elif style.startswith("Heading 3"):
                lines.append(f"### {text}")
            elif style.startswith("List"):
                lines.append(f"- {text}")
            else:
                lines.append(text)
        # Tables
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            header = "| " + " | ".join(rows[0]) + " |"
            sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
            lines.extend([header, sep])
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)
    except ImportError:
        return ""
    except Exception:
        return ""


def xlsx_to_markdown(path: Path) -> str:
    """Convert an .xlsx file to markdown text using openpyxl."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        sections = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                # Skip entirely empty rows
                if all(cell is None for cell in row):
                    continue
                rows.append([str(cell) if cell is not None else "" for cell in row])
            if not rows:
                continue
            sections.append(f"## Sheet: {sheet_name}")
            if len(rows) >= 1:
                header = "| " + " | ".join(rows[0]) + " |"
                sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
                sections.extend([header, sep])
                for row in rows[1:]:
                    sections.append("| " + " | ".join(row) + " |")
        wb.close()
        return "\n".join(sections)
    except ImportError:
        return ""
    except Exception:
        return ""


def convert_office_file(path: Path, out_dir: Path) -> Path | None:
    """Convert a .docx or .xlsx to a markdown sidecar in out_dir.

    Returns the path of the converted .md file, or None if conversion failed
    or the required library is not installed.
    """
    ext = path.suffix.lower()
    if ext == ".docx":
        text = docx_to_markdown(path)
    elif ext == ".xlsx":
        text = xlsx_to_markdown(path)
    else:
        return None

    if not text.strip():
        return None

    out_dir.mkdir(parents=True, exist_ok=True)
    # Use a stable name derived from the original path to avoid collisions
    import hashlib
    name_hash = hashlib.sha256(str(path.resolve()).encode()).hexdigest()[:8]
    out_path = out_dir / f"{path.stem}_{name_hash}.md"
    out_path.write_text(
        f"<!-- converted from {path.name} -->\n\n{text}",
        encoding="utf-8",
    )
    return out_path


def count_words(path: Path) -> int:
    try:
        ext = path.suffix.lower()
        if ext == ".pdf":
            return len(extract_pdf_text(path).split())
        if ext == ".docx":
            return len(docx_to_markdown(path).split())
        if ext == ".xlsx":
            return len(xlsx_to_markdown(path).split())
        return len(path.read_text(encoding="utf-8", errors="ignore").split())
    except Exception:
        return 0


# graphify's own output directory — always pruned by default to prevent
# self-ingestion loops (e.g. re-running --obsidian from a vault root would
# otherwise re-ingest prior exported notes as fresh document inputs).
# The graphify-out/memory/ subtree is RE-included explicitly in detect()
# via the scan_paths allow-list below.
_SELF_OUTPUT_DIRS = {"graphify-out", "graphify_out"}

_OUTPUT_MANIFEST_MAX_RUNS = 5

# Large generated files that are never useful to extract
_SKIP_FILES = {
    "package-lock.json", "yarn.lock", "pnpm-lock.yaml",
    "Cargo.lock", "poetry.lock", "Gemfile.lock",
    "composer.lock", "go.sum", "go.work.sum",
}


def _load_graphifyignore(root: Path) -> list[str]:
    """Read .graphifyignore from root **and ancestor directories**, returning patterns.

    Walks upward from *root* towards the filesystem root, collecting patterns
    from every ``.graphifyignore`` encountered (like ``.gitignore`` discovery).
    The search stops at the filesystem root or at a ``.git`` directory boundary
    so it doesn't leak outside the repository.

    Lines starting with # are comments. Blank lines are ignored.
    Patterns follow gitignore semantics: glob matched against the path
    relative to root. A leading slash anchors to root. A trailing slash
    matches directories only (we match both dir and file for simplicity).
    """
    patterns: list[str] = []
    current = root.resolve()
    while True:
        ignore_file = current / ".graphifyignore"
        if ignore_file.exists():
            for line in ignore_file.read_text(encoding="utf-8", errors="ignore").splitlines():
                line = line.strip()
                if line and not line.startswith("#"):
                    patterns.append(line)
        # Stop climbing once we've processed the git repo root
        if (current / ".git").exists():
            break
        parent = current.parent
        if parent == current:
            break  # filesystem root
        current = parent
    return patterns


def _is_ignored(path: Path, root: Path, patterns: list[str]) -> bool:
    """Return True if path matches any .graphifyignore pattern."""
    if not patterns:
        return False
    try:
        rel = str(path.relative_to(root))
    except ValueError:
        return False
    rel = rel.replace(os.sep, "/")
    parts = rel.split("/")
    for pattern in patterns:
        # Normalize: strip leading/trailing slashes for matching purposes
        p = pattern.strip("/")
        if not p:
            continue
        # Match against full relative path
        if fnmatch.fnmatch(rel, p):
            return True
        # Match against filename alone
        if fnmatch.fnmatch(path.name, p):
            return True
        # Match against any path segment or prefix
        # e.g. "vendor" or "vendor/" should match "vendor/lib.py"
        for i, part in enumerate(parts):
            if fnmatch.fnmatch(part, p):
                return True
            if fnmatch.fnmatch("/".join(parts[:i + 1]), p):
                return True
    return False


def _save_output_manifest(
    artifacts_dir: Path,
    notes_dir: Path,
    written_files: list[str],
    run_id: str | None = None,
) -> None:
    """Append a run entry and write output-manifest.json atomically (D-29).

    - GC's stale file entries from prior runs (D-28: removes entries where path no longer exists)
    - Appends new run entry with run_id, timestamp, notes_dir, artifacts_dir, files
    - FIFO-trims runs list to N=5 (D-24)
    - Writes atomically via tmp + fsync + os.replace; cleans tmp on OSError; re-raises
    """
    manifest_path = artifacts_dir / _OUTPUT_MANIFEST_NAME
    manifest_path.parent.mkdir(parents=True, exist_ok=True)

    existing = _load_output_manifest(artifacts_dir)
    ts = datetime.datetime.now(datetime.timezone.utc).isoformat()
    if run_id is None:
        h = hashlib.sha256(f"{notes_dir}{ts}".encode()).hexdigest()[:8]
        run_id = f"{ts}-{h}"

    new_run = {
        "run_id": run_id,
        "timestamp": ts,
        "notes_dir": str(notes_dir.resolve()),
        "artifacts_dir": str(artifacts_dir.resolve()),
        "files": [str(Path(f).resolve()) for f in written_files],
    }

    runs: list[dict] = existing.get("runs", [])
    # D-28: GC stale file entries from prior runs
    for run in runs:
        run["files"] = [f for f in run.get("files", []) if Path(f).exists()]
    # Append and FIFO-trim to N=5 (D-24)
    runs.append(new_run)
    runs = runs[-_OUTPUT_MANIFEST_MAX_RUNS:]

    manifest = {"version": _OUTPUT_MANIFEST_VERSION, "runs": runs}

    # Atomic write: tmp + os.replace (D-29) — mirrors merge.py:_write_atomic
    tmp = manifest_path.with_suffix(".json.tmp")
    try:
        with open(tmp, "w", encoding="utf-8") as fh:
            fh.write(json.dumps(manifest, indent=2, sort_keys=True))
            fh.flush()
            os.fsync(fh.fileno())
        os.replace(tmp, manifest_path)
    except OSError:
        if tmp.exists():
            try:
                tmp.unlink()
            except OSError:
                pass
        raise


def _relative_posix_under_root(path: Path, root: Path) -> str | None:
    try:
        return path.resolve().relative_to(root.resolve()).as_posix()
    except ValueError:
        return None


def _dot_graphify_from_profile(profile: dict | None) -> dict:
    defaults = {
        "include_globs": [],
        "exclude_globs": ["**/*.yaml", "**/profile.yaml"],
        "auto_track_discoveries": False,
        "tracked_paths": [],
    }
    if not profile:
        return dict(defaults)
    corpus = profile.get("corpus")
    if not isinstance(corpus, dict):
        return dict(defaults)
    dot = corpus.get("dot_graphify")
    if not isinstance(dot, dict):
        return dict(defaults)
    out = dict(defaults)
    for k in defaults:
        if k in dot:
            out[k] = dot[k]
    return out


def _dot_graphify_path_eligible(rel_posix: str, dot: dict) -> bool:
    if not rel_posix.startswith(".graphify/"):
        return True
    low = rel_posix.lower()
    if low.endswith(".yaml") or low.endswith(".yml"):
        return False
    if rel_posix.split("/")[-1] == "profile.yaml":
        return False
    inc = dot.get("include_globs") or []
    exc = dot.get("exclude_globs") or []
    if not inc:
        return False
    if not any(fnmatch.fnmatch(rel_posix, g) for g in inc):
        return False
    if any(fnmatch.fnmatch(rel_posix, g) for g in exc):
        return False
    return True


def detect(
    root: Path,
    *,
    follow_symlinks: bool = False,
    resolved: "ResolvedOutput | None" = None,
    profile: dict | None = None,
) -> dict:
    files: dict[FileType, list[str]] = {
        FileType.CODE: [],
        FileType.DOCUMENT: [],
        FileType.PAPER: [],
        FileType.IMAGE: [],
        FileType.VIDEO: [],
    }
    total_words = 0

    skipped_sensitive: list[str] = []
    ignore_patterns = _load_graphifyignore(root)

    # Plan 29-02 (D-39): single source of truth for skip decisions, surfaced
    # via additive return key `skipped`. Each pruning site below also appends
    # to skipped[reason]. T-29-05: cap each list at 10000 entries to bound
    # memory on pathological inputs; overflow counts go to skipped_overflow.
    _SKIP_CAP = 10000
    skipped: dict[str, list[str]] = {
        "nesting": [],
        "exclude-glob": [],
        "manifest": [],
        "sensitive": [],
        "noise-dir": [],
    }
    skipped_overflow: dict[str, int] = {k: 0 for k in skipped}

    def _record_skip(reason: str, rel_path: str) -> None:
        if len(skipped[reason]) < _SKIP_CAP:
            skipped[reason].append(rel_path)
        else:
            skipped_overflow[reason] += 1

    # Phase 28: compute resolved-aware basenames and combined exclude patterns
    resolved_basenames: frozenset[str] = frozenset()
    if resolved is not None:
        resolved_basenames = frozenset({
            resolved.notes_dir.name,
            resolved.artifacts_dir.name,
        }) - _SELF_OUTPUT_DIRS

    exclude_globs: list[str] = list(resolved.exclude_globs) if resolved else []
    all_ignore_patterns: list[str] = list(ignore_patterns) + exclude_globs

    nested_paths: list[str] = []

    # Phase 45 (D-45.02/D-45.03): manifest paths from resolved artifacts_dir and/or default graphify-out.
    prior_files = build_prior_files(root, resolved)

    # Always include graphify-out/memory/ - query results filed back into the graph
    memory_dir = root / "graphify-out" / "memory"
    scan_paths = [root]
    if memory_dir.exists():
        scan_paths.append(memory_dir)

    seen: set[Path] = set()
    all_files: list[Path] = []

    for scan_root in scan_paths:
        in_memory_tree = memory_dir.exists() and str(scan_root).startswith(str(memory_dir))
        for dirpath, dirnames, filenames in os.walk(scan_root, followlinks=follow_symlinks):
            dp = Path(dirpath)
            if follow_symlinks and os.path.islink(dirpath):
                real = os.path.realpath(dirpath)
                parent_real = os.path.realpath(os.path.dirname(dirpath))
                if parent_real == real or parent_real.startswith(real + os.sep):
                    dirnames.clear()
                    continue
            if not in_memory_tree:
                # Prune noise dirs in-place so os.walk never descends into them.
                # Accumulate nesting paths separately for the D-20 single-line warning.
                pruned: set[str] = set()
                for d in dirnames:
                    reason = dir_prune_reason(
                        d,
                        dp,
                        root,
                        resolved_basenames=resolved_basenames,
                        patterns=all_ignore_patterns,
                    )
                    if reason:
                        _record_skip(reason, str(dp / d))
                        if reason == "nesting":
                            nested_paths.append(str(dp / d))
                        pruned.add(d)
                dirnames[:] = [d for d in dirnames if d not in pruned]
            for fname in filenames:
                if fname in _SKIP_FILES:
                    continue
                p = dp / fname
                if p not in seen:
                    seen.add(p)
                    all_files.append(p)

    # D-20: emit a single summary warning if any nesting paths were pruned
    if nested_paths:
        deepest = max(nested_paths, key=lambda p: p.count(os.sep))
        print(
            f"[graphify] WARNING: skipped {len(nested_paths)} nested output path(s) "
            f"(deepest: {deepest})",
            file=sys.stderr,
        )

    converted_dir = root / "graphify-out" / "converted"

    dot_cfg = _dot_graphify_from_profile(profile)
    dot_discovered: list[str] = []

    for p in all_files:
        # For memory dir files, skip hidden/noise filtering
        in_memory = memory_dir.exists() and str(p).startswith(str(memory_dir))
        if not in_memory:
            # Hidden files are already excluded via dir pruning above,
            # but catch hidden files at the root level
            if p.name.startswith("."):
                continue
            # Skip files inside our own converted/ dir (avoid re-processing sidecars)
            if str(p).startswith(str(converted_dir)):
                continue
        rel_under = _relative_posix_under_root(p, root)
        if rel_under is not None and rel_under.startswith(".graphify/"):
            if not _dot_graphify_path_eligible(rel_under, dot_cfg):
                continue
            dot_discovered.append(rel_under)
        if _is_ignored(p, root, all_ignore_patterns):
            _record_skip("exclude-glob", str(p))
            continue
        # Phase 28 D-27: silent skip for files recorded in a prior output-manifest run
        if prior_files and str(p.resolve()) in prior_files:
            _record_skip("manifest", str(p))
            continue
        if _is_sensitive(p):
            skipped_sensitive.append(str(p))
            _record_skip("sensitive", str(p))
            continue
        ftype = classify_file(p)
        if ftype:
            # Office files: convert to markdown sidecar so subagents can read them
            if p.suffix.lower() in OFFICE_EXTENSIONS:
                md_path = convert_office_file(p, converted_dir)
                if md_path:
                    files[ftype].append(str(md_path))
                    total_words += count_words(md_path)
                else:
                    # Conversion failed (library not installed) - skip with note
                    skipped_sensitive.append(str(p) + " [office conversion failed - pip install graphifyy[office]]")
                continue
            files[ftype].append(str(p))
            if ftype != FileType.VIDEO:
                total_words += count_words(p)

    if skipped["manifest"]:
        n = len(skipped["manifest"]) + skipped_overflow["manifest"]
        print(
            f"[graphify] skipped {n} prior-output file(s) per output-manifest.json "
            "— see graphify doctor --dry-run for detail",
            file=sys.stderr,
        )

    total_files = sum(len(v) for v in files.values())
    needs_graph = total_words >= CORPUS_WARN_THRESHOLD

    # Determine warning - lower bound, upper bound, or sensitive files skipped
    warning: str | None = None
    if not needs_graph:
        warning = (
            f"Corpus is ~{total_words:,} words - fits in a single context window. "
            f"You may not need a graph."
        )
    elif total_words >= CORPUS_UPPER_THRESHOLD or total_files >= FILE_COUNT_UPPER:
        warning = (
            f"Large corpus: {total_files} files · ~{total_words:,} words. "
            f"Semantic extraction will be expensive (many Claude tokens). "
            f"Consider running on a subfolder, or use --no-semantic to run AST-only."
        )

    return {
        "files": {k.value: v for k, v in files.items()},
        "total_files": total_files,
        "total_words": total_words,
        "needs_graph": needs_graph,
        "warning": warning,
        "skipped_sensitive": skipped_sensitive,
        "graphifyignore_patterns": len(ignore_patterns),
        "skipped": skipped,
        "dot_graphify_discovered": sorted(set(dot_discovered)),
    }


def load_manifest(manifest_path: str = _MANIFEST_PATH) -> dict[str, float]:
    """Load the file modification time manifest from a previous run."""
    try:
        return json.loads(Path(manifest_path).read_text(encoding="utf-8"))
    except Exception:
        return {}


def save_manifest(files: dict[str, list[str]], manifest_path: str = _MANIFEST_PATH) -> None:
    """Save current file mtimes so the next --update run can diff against them.

    Written atomically (tmp + fsync + os.replace) like ``_save_output_manifest`` so
    concurrent runs or crashes cannot leave a truncated ``manifest.json``.
    """
    manifest: dict[str, float] = {}
    for file_list in files.values():
        for f in file_list:
            try:
                manifest[f] = Path(f).stat().st_mtime
            except OSError:
                pass  # file deleted between detect() and manifest write - skip it
    dest = Path(manifest_path)
    dest.parent.mkdir(parents=True, exist_ok=True)
    payload = json.dumps(manifest, indent=2, sort_keys=True)
    tmp = dest.with_suffix(".json.tmp")
    try:
        with open(tmp, "w", encoding="utf-8") as fh:
            fh.write(payload)
            fh.flush()
            os.fsync(fh.fileno())
        os.replace(tmp, dest)
    except OSError:
        if tmp.exists():
            try:
                tmp.unlink()
            except OSError:
                pass
        raise


def detect_incremental(root: Path, manifest_path: str = _MANIFEST_PATH) -> dict:
    """Like detect(), but returns only new or modified files since the last run.

    Compares current file mtimes against the stored manifest.
    Use for --update mode: re-extract only what changed, merge into existing graph.
    """
    full = detect(root)
    manifest = load_manifest(manifest_path)

    if not manifest:
        # No previous run - treat everything as new
        full["incremental"] = True
        full["new_files"] = full["files"]
        full["unchanged_files"] = {k: [] for k in full["files"]}
        full["new_total"] = full["total_files"]
        return full

    new_files: dict[str, list[str]] = {k: [] for k in full["files"]}
    unchanged_files: dict[str, list[str]] = {k: [] for k in full["files"]}

    for ftype, file_list in full["files"].items():
        for f in file_list:
            stored_mtime = manifest.get(f)
            try:
                current_mtime = Path(f).stat().st_mtime
            except Exception:
                current_mtime = 0
            if stored_mtime is None or current_mtime > stored_mtime:
                new_files[ftype].append(f)
            else:
                unchanged_files[ftype].append(f)

    # Files in manifest that no longer exist - their cached nodes are now ghost nodes
    current_files = {f for flist in full["files"].values() for f in flist}
    deleted_files = [f for f in manifest if f not in current_files]

    new_total = sum(len(v) for v in new_files.values())
    full["incremental"] = True
    full["new_files"] = new_files
    full["unchanged_files"] = unchanged_files
    full["new_total"] = new_total
    full["deleted_files"] = deleted_files
    return full
